"""
Word 文档加载器
优先用 python-docx 直接提取文字结构
嵌入图片/表格图片走 VLM
"""

import io
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image


def _get_paragraph_level(para) -> int:
    """从段落样式推断标题级别，0=正文"""
    style_name = para.style.name.lower()
    if "heading 1" in style_name or "标题 1" in style_name:
        return 1
    if "heading 2" in style_name or "标题 2" in style_name:
        return 2
    if "heading 3" in style_name or "标题 3" in style_name:
        return 3
    if "heading 4" in style_name or "标题 4" in style_name:
        return 4
    if "heading 5" in style_name or "标题 5" in style_name:
        return 5
    if "heading 6" in style_name or "标题 6" in style_name:
        return 6
    # fallback: 通过字号推断
    for run in para.runs:
        if run.font.size and run.font.size >= Pt(16):
            return 2
        if run.font.size and run.font.size >= Pt(14):
            return 3
    return 0


def _extract_paragraph_images(para, out_dir: Path, img_counter: list) -> list[dict]:
    """提取段落内嵌图片，返回图片块列表"""
    images = []
    for elem in para._element.iter():
        if elem.tag.endswith("}blip"):
            rId = elem.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if rId:
                img_counter[0] += 1
                images.append({"_rId": rId, "_counter": img_counter[0]})
    return images


def _table_to_text_grid(table) -> list[list[str]]:
    """将 docx Table 转为二维字符串数组"""
    grid = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        grid.append(row_data)
    return grid


def _grid_to_html(grid: list[list[str]]) -> str:
    """二维数组转标准 HTML 表格"""
    if not grid:
        return ""
    lines = ['<table border="1">']
    for ri, row in enumerate(grid):
        lines.append("  <tr>")
        tag = "th" if ri == 0 else "td"
        for cell in row:
            cell_esc = cell.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            lines.append(f"    <{tag}>{cell_esc}</{tag}>")
        lines.append("  </tr>")
    lines.append("</table>")
    return "\n".join(lines)

def _grid_to_markdown(grid):
    """兼容旧调用"""
    return _grid_to_html(grid)


def _grid_to_png(grid: list[list[str]], out_path) -> bool:
    """
    用 Pillow 把二维字符串数组渲染成 PNG 表格图片。
    单元格内容超长时自动换行，不截断。
    """
    try:
        import os
        from PIL import Image, ImageDraw, ImageFont

        if not grid:
            return False

        # ── 跨平台中文字体候选 ────────────────────────
        FONT_CANDIDATES = [
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\msyhbd.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\simsun.ttc",
            r"C:\Windows\Fonts\simfang.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Light.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
        ]

        def _find_font(size):
            for path in FONT_CANDIDATES:
                if os.path.isfile(path):
                    try:
                        return ImageFont.truetype(path, size)
                    except Exception:
                        pass
            return ImageFont.load_default()

        FONT_SIZE   = 12
        HEADER_ROWS = 2
        PAD_X       = 6
        PAD_Y       = 4
        BORDER      = 1
        MAX_COL_W   = 160   # 每列最大宽度（像素），超过则换行
        LINE_GAP    = 3     # 行间距

        font_hdr  = _find_font(FONT_SIZE)
        font_body = _find_font(FONT_SIZE)

        dummy = Image.new("RGB", (1, 1))
        dc    = ImageDraw.Draw(dummy)

        def measure_text(text, fnt):
            try:
                bb = dc.textbbox((0, 0), text, font=fnt)
                return bb[2] - bb[0], bb[3] - bb[1]
            except Exception:
                return len(text) * (FONT_SIZE // 2 + 1), FONT_SIZE

        def char_w(fnt):
            """估算平均字符宽度"""
            w, _ = measure_text("国", fnt)
            return max(w, 1)

        def wrap_text(text, fnt, max_w):
            """把文本按 max_w 像素宽度换行，返回行列表"""
            if not text:
                return [""]
            cw = char_w(fnt)
            max_chars_per_line = max(1, int(max_w / cw))
            lines = []
            while text:
                # 二分查找能放下的最大字符数
                lo, hi = 1, min(len(text), max_chars_per_line * 2)
                while lo < hi:
                    mid = (lo + hi + 1) // 2
                    w, _ = measure_text(text[:mid], fnt)
                    if w <= max_w:
                        lo = mid
                    else:
                        hi = mid - 1
                lines.append(text[:lo])
                text = text[lo:]
            return lines if lines else [""]

        ncols = max(len(row) for row in grid)
        nrows = len(grid)

        # ── 第一遍：确定每列宽度（先按内容，限制最大值）─────
        col_content_w = [0] * ncols
        for ri, row in enumerate(grid):
            fnt = font_hdr if ri < HEADER_ROWS else font_body
            for ci in range(ncols):
                val = str(row[ci]).strip() if ci < len(row) else ""
                w, _ = measure_text(val, fnt)
                col_content_w[ci] = max(col_content_w[ci], w + PAD_X * 2)

        col_widths = [min(cw, MAX_COL_W + PAD_X * 2) for cw in col_content_w]

        # ── 第二遍：确定每行高度（按换行后的实际行数）────────
        line_h = FONT_SIZE + LINE_GAP
        cell_lines = []   # cell_lines[ri][ci] = list of text lines

        for ri, row in enumerate(grid):
            fnt = font_hdr if ri < HEADER_ROWS else font_body
            row_cell_lines = []
            max_line_count = 1
            for ci in range(ncols):
                val     = str(row[ci]).strip() if ci < len(row) else ""
                inner_w = col_widths[ci] - PAD_X * 2
                lines   = wrap_text(val, fnt, inner_w)
                row_cell_lines.append(lines)
                max_line_count = max(max_line_count, len(lines))
            cell_lines.append(row_cell_lines)
            # 行高 = 最多行数 * 行高 + 上下留白
            # (header行至少保证单行高度)

        row_heights = []
        for ri in range(nrows):
            max_lines = max(len(cell_lines[ri][ci]) for ci in range(ncols))
            rh = max_lines * line_h + PAD_Y * 2
            row_heights.append(rh)

        total_w = sum(col_widths)  + BORDER * (ncols + 1)
        total_h = sum(row_heights) + BORDER * (nrows + 1)

        img = Image.new("RGB", (total_w, total_h), "white")
        d   = ImageDraw.Draw(img)

        HEADER_BG  = (31,  56, 100)
        HEADER_FG  = (255, 255, 255)
        ALT_BG     = (235, 240, 248)
        NORMAL_BG  = (255, 255, 255)
        BODY_FG    = (26,  26,  26)
        BORDER_COL = (136, 136, 136)

        y = BORDER
        for ri in range(nrows):
            x      = BORDER
            fnt    = font_hdr if ri < HEADER_ROWS else font_body
            is_hdr = ri < HEADER_ROWS
            row_bg = HEADER_BG if is_hdr else (ALT_BG if ri % 2 == 0 else NORMAL_BG)
            row_fg = HEADER_FG if is_hdr else BODY_FG
            rh     = row_heights[ri]

            for ci in range(ncols):
                cw    = col_widths[ci]
                lines = cell_lines[ri][ci]

                d.rectangle([x, y, x + cw, y + rh], fill=row_bg, outline=BORDER_COL)

                # 多行文字，垂直居中整体
                total_text_h = len(lines) * line_h - LINE_GAP
                ty_start = y + (rh - total_text_h) // 2
                for li, line in enumerate(lines):
                    lw, lh = measure_text(line, fnt)
                    tx = x + (cw - lw) // 2
                    ty = ty_start + li * line_h
                    d.text((tx, ty), line, fill=row_fg, font=fnt)

                x += cw + BORDER
            y += rh + BORDER

        img.save(str(out_path), "PNG", dpi=(150, 150))
        return True

    except Exception:
        import traceback
        traceback.print_exc()
        return False


def load_word(file_path: str, out_dir: Path, progress_cb=None) -> dict:
    """
    加载 Word 文档，提取结构化内容块
    out_dir: 图片保存目录
    返回 {source, doc_type, blocks}
    blocks 中每个元素：
      {type, level, text, md_table, path, filename, rId, needs_vlm}
    """
    (out_dir / "images").mkdir(parents=True, exist_ok=True)
    img_dir = out_dir / "images"

    doc = Document(file_path)
    blocks = []
    img_counter = [0]
    table_counter = [0]
    total_elements = len(doc.paragraphs) + len(doc.tables)

    # 遍历文档体中的所有元素（保持顺序）
    body = doc.element.body
    processed = 0

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        # ── 段落 ─────────────────────────────────────
        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            text = para.text.strip()

            # 检查是否包含图片
            inline_imgs = _extract_paragraph_images(para, img_dir, img_counter)
            if inline_imgs:
                for img_info in inline_imgs:
                    rId = img_info["_rId"]
                    counter = img_info["_counter"]
                    fname = f"fig_{counter}.png"
                    fpath = img_dir / fname
                    try:
                        rel = doc.part.rels.get(rId)
                        if rel:
                            img_bytes = rel.target_part.blob
                            pil_img = Image.open(io.BytesIO(img_bytes))
                            pil_img.save(str(fpath), "PNG")
                            blocks.append({
                                "type": "figure",
                                "level": 0,
                                "text": "",
                                "path": str(fpath),
                                "filename": fname,
                                "needs_vlm": False,
                            })
                    except Exception:
                        pass

            if text:
                level = _get_paragraph_level(para)
                blocks.append({
                    "type": "heading" if level > 0 else "text",
                    "level": level,
                    "text": text,
                    "path": None,
                    "filename": None,
                    "needs_vlm": False,
                })

        # ── 表格 ─────────────────────────────────────
        elif tag == "tbl":
            from docx.table import Table
            table_counter[0] += 1
            tbl = Table(child, doc)
            grid = _table_to_text_grid(tbl)
            md_table = _grid_to_markdown(grid)  # 保留，供 tool3 QA 用

            # 渲染为 PNG 图片
            tables_dir = out_dir / "tables"
            tables_dir.mkdir(parents=True, exist_ok=True)
            png_fname = f"table_{table_counter[0]:03d}.png"
            png_path = tables_dir / png_fname
            png_ok = _grid_to_png(grid, png_path)

            blocks.append({
                "type": "table",
                "level": 0,
                "text": "",
                "md_table": md_table,       # HTML格式，供 tool3 QA 生成
                "grid": grid,               # 原始二维数组，供 tool3 QA 生成
                "path": str(png_path) if png_ok else None,
                "filename": png_fname if png_ok else None,
                "needs_vlm": False,
            })

        processed += 1
        if progress_cb and total_elements > 0:
            progress_cb(processed, total_elements,
                        f"解析 Word 内容 {processed}/{total_elements}...")

    return {
        "source": file_path,
        "doc_type": "word",
        "blocks": blocks,
        "out_dir": str(out_dir),
    }
