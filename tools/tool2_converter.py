"""
Tool 2 — MD → Word 转换面板
完整移植自 mineru_converter2_1.py，功能无损
"""
import os
import sys
import re
import threading
import traceback
from pathlib import Path

import tkinter as tk
from tkinter import filedialog, messagebox, ttk

from gui.theme import COLORS, FONTS, PADDING
from gui.widgets import (
    StyledButton, FilePickRow, LogView, ProgressRow,
    SectionHeader, Divider
)


# ══════════════════════════════════════════════════════════════
#  核心转换逻辑（完整从 mineru_converter2_1.py 移植）
# ══════════════════════════════════════════════════════════════

def parse_markdown(text: str, md_dir: str = "") -> list:
    """
    解析 annotated.md，生成 token 列表。
    支持：
      - #### 分段符（单独成行，转Word时输出为分隔段落）
      - @@@tag1@@@tag2@@@ 标签行（单独输出为灰色小字段落）
      - tags: @@@...@@@ 旧格式标签行（同上）
      - <!-- TABLE:tables/table_N.html --> 表格占位符（读HTML→Word表格）
      - <table>...</table> 内联HTML表格
      - | pipe | 表格 → Word表格
      - Q:...;A:... 行（表格QA对，正常段落输出）
      - 其余正常MD语法
    """
    tokens = []
    lines = text.splitlines()
    i = 0
    para_buf = []

    def flush_para(buf):
        t = " ".join(l.strip() for l in buf if l.strip())
        if t:
            tokens.append({"type": "paragraph", "text": t})

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── #### 分段符 → 细分隔线 ──────────────────
        if stripped == "####":
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "segment_sep"})
            i += 1; continue

        # ── @@@tags@@@ 标签行（新格式）→ 原样写入 ──
        if stripped.startswith("@@@") and stripped.endswith("@@@") and len(stripped) > 6:
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "tags_raw", "text": stripped})
            i += 1; continue

        # ── tags: @@@...@@@ 旧格式 → 原样写入 ───────
        if re.match(r"^tags:\s*@@@", stripped):
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "tags_raw", "text": stripped})
            i += 1; continue

        # ── <!-- TABLE:tables/xxx.png --> 占位符 ─────
        m = re.match(r"^<!--\s*TABLE:(.*?)\s*-->$", stripped)
        if m:
            flush_para(para_buf); para_buf = []
            tbl_path = m.group(1).strip()
            # 跳过随后的 ![表格](...) 行（避免重复输出图片）
            next_i = i + 1
            if next_i < len(lines) and re.match(r"^!\[.*\]\(.*\)", lines[next_i].strip()):
                i = next_i  # 消耗掉那行
            tokens.append({"type": "table_image", "path": tbl_path})
            i += 1; continue

        # ── 其他 HTML 注释跳过 ───────────────────────
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            flush_para(para_buf); para_buf = []
            i += 1; continue

        # ── 代码块 ──────────────────────────────────────────
        if line.startswith("```"):
            flush_para(para_buf); para_buf = []
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i]); i += 1
            tokens.append({"type": "code_block", "text": "\n".join(code_lines)})
            i += 1; continue

        # ── HTML 表格 ────────────────────────────────────────
        if re.search(r"<table", stripped, re.IGNORECASE) and not stripped.startswith("#"):
            flush_para(para_buf); para_buf = []
            html_lines = [line]
            if re.search(r"</table>", line, re.IGNORECASE):
                i += 1
            else:
                i += 1
                while i < len(lines):
                    html_lines.append(lines[i])
                    if re.search(r"</table>", lines[i], re.IGNORECASE):
                        i += 1; break
                    i += 1
            tokens.append({"type": "html_table", "html": "\n".join(html_lines)})
            continue

        # ── Pipe 表格 → 转为 html_table token ───────────────
        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_para(para_buf); para_buf = []
            pipe_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                pipe_lines.append(lines[i])
                i += 1
            html = _pipe_to_html(pipe_lines)
            if html:
                tokens.append({"type": "html_table", "html": html})
            continue

        # ── 标题 ─────────────────────────────────────────────
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "heading", "level": len(m.group(1)),
                           "text": m.group(2).strip()})
            i += 1; continue

        # ── 水平线 ───────────────────────────────────────────
        if re.match(r"^\s*[-*_]{3,}\s*$", line):
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "hr"}); i += 1; continue

        # ── 独立图片行 ───────────────────────────────────────
        m = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if m:
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "image", "alt": m.group(1), "src": m.group(2)})
            i += 1; continue

        # ── 列表项 ───────────────────────────────────────────
        m = re.match(r"^(\s*)[*\-+]\s+(.*)", line)
        if m:
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "list_item", "ordered": False,
                           "level": len(m.group(1)) // 2, "text": m.group(2).strip()})
            i += 1; continue

        m = re.match(r"^(\s*)\d+[.)]\s+(.*)", line)
        if m:
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "list_item", "ordered": True,
                           "level": len(m.group(1)) // 2, "text": m.group(2).strip()})
            i += 1; continue

        # ── 空行 ─────────────────────────────────────────────
        if stripped == "":
            flush_para(para_buf); para_buf = []
            i += 1; continue

        # ── 行内图片 ─────────────────────────────────────────
        if "![" in line:
            flush_para(para_buf); para_buf = []
            remaining = line
            while remaining:
                bm = re.match(r"(.*?)!\[([^\]]*)\]\(([^)]+)\)(.*)", remaining, re.DOTALL)
                if bm:
                    if bm.group(1).strip():
                        tokens.append({"type": "paragraph", "text": bm.group(1).strip()})
                    tokens.append({"type": "image", "alt": bm.group(2), "src": bm.group(3)})
                    remaining = bm.group(4)
                else:
                    if remaining.strip():
                        tokens.append({"type": "paragraph", "text": remaining.strip()})
                    break
            i += 1; continue

        # ── Q: 问题行 → 单独段落，不合并 ──────────────
        if stripped.startswith("Q:"):
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "qa_line", "text": stripped})
            i += 1; continue

        # ── A: 回答行 → 单独段落，不合并 ──────────────
        if stripped.startswith("A:"):
            flush_para(para_buf); para_buf = []
            tokens.append({"type": "qa_line", "text": stripped})
            i += 1; continue

        # ── 普通文本（含 Q:...;A:... 的QA行） ───────────────
        para_buf.append(line)
        i += 1

    flush_para(para_buf)
    return tokens


def _pipe_to_html(pipe_lines: list) -> str:
    """把 pipe 表格行转成 HTML 表格"""
    rows = []
    is_header_row = []
    for line in pipe_lines:
        s = line.strip()
        if not s or not s.startswith("|"):
            continue
        # 分隔行（|---|---| 格式）跳过
        if re.match(r"^\|[\s\-:]+\|", s):
            continue
        cells = [c.strip() for c in s.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return ""

    lines = ['<table border="1">']
    for ri, row in enumerate(rows):
        lines.append("  <tr>")
        tag = "th" if ri == 0 else "td"
        for cell in row:
            esc = cell.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            lines.append(f"    <{tag}>{esc}</{tag}>")
        lines.append("  </tr>")
    lines.append("</table>")
    return "\n".join(lines)


def parse_html_table(html: str) -> dict:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        table = soup.find("table")
        if not table:
            return {"headers": [], "rows": [], "col_count": 0}
        headers, rows = [], []
        for tr in table.find_all("tr"):
            cells = tr.find_all(["th", "td"])
            data = [c.get_text(separator=" ", strip=True) for c in cells]
            if cells and all(c.name == "th" for c in cells):
                headers = data
            else:
                rows.append(data)
        col_count = max(len(headers), max((len(r) for r in rows), default=0)) if (headers or rows) else 0
        return {"headers": headers, "rows": rows, "col_count": col_count}
    except ImportError:
        return {"headers": [], "rows": [], "col_count": 0}


def resolve_image(src: str, images_dir: str, md_dir: str = ""):
    candidates = [
        src,
        os.path.join(images_dir, src),
        os.path.join(images_dir, os.path.basename(src)),
    ]
    if md_dir:
        candidates.append(os.path.join(md_dir, src))
        candidates.append(os.path.join(md_dir, os.path.basename(src)))
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def add_inline(paragraph, text: str):
    try:
        from docx.shared import Pt
    except ImportError:
        paragraph.add_run(text)
        return

    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"
        r"|(\*\*(.+?)\*\*)"
        r"|(__(.+?)__)"
        r"|(\*(.+?)\*)"
        r"|(_(.+?)_)"
        r"|(`(.+?)`)"
    )
    last = 0
    for m in pattern.finditer(text):
        s, e = m.span()
        if s > last:
            paragraph.add_run(text[last:s])
        if m.group(1):
            r = paragraph.add_run(m.group(2)); r.italic = True
        elif m.group(3):
            paragraph.add_run(m.group(4))
        elif m.group(5):
            paragraph.add_run(m.group(6))
        elif m.group(7):
            r = paragraph.add_run(m.group(8)); r.italic = True
        elif m.group(9):
            r = paragraph.add_run(m.group(10)); r.italic = True
        elif m.group(11):
            r = paragraph.add_run(m.group(12))
            r.font.name = "Courier New"; r.font.size = Pt(9)
        last = e
    if last < len(text):
        paragraph.add_run(text[last:])


def build_docx_file(tokens: list, images_dir: str, output_path: str,
                    log_cb=None, md_dir: str = ""):
    def log(msg):
        if log_cb: log_cb(msg)

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise RuntimeError(f"缺少依赖 python-docx: {e}")

    HEADING_COLOR = {
        1: RGBColor(0x1F, 0x38, 0x64),
        2: RGBColor(0x2E, 0x50, 0x90),
        3: RGBColor(0x36, 0x5F, 0x91),
        4: RGBColor(0x44, 0x72, 0xC4),
    }
    CONTENT_WIDTH = Cm(16.5)

    def _set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width   = Cm(21)
    sec.page_height  = Cm(29.7)
    sec.left_margin  = sec.right_margin  = Cm(2.54)
    sec.top_margin   = sec.bottom_margin = Cm(2.54)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    for lvl, size in [(1,20),(2,16),(3,14),(4,12)]:
        try:
            st = doc.styles[f"Heading {lvl}"]
            st.font.name = "Arial"
            st.font.size = Pt(size)
            st.font.bold = False
            if lvl in HEADING_COLOR:
                st.font.color.rgb = HEADING_COLOR[lvl]
        except Exception:
            pass

    for tok in tokens:
        t = tok["type"]

        if t == "heading":
            lvl = min(tok["level"], 6)
            p = doc.add_heading(level=lvl)
            p.clear()
            run = p.add_run(tok["text"])
            run.bold = None
            run.font.name = "Arial"
            run.font.size = Pt({1:20,2:16,3:14,4:12,5:11,6:10}.get(lvl, 11))
            if lvl in HEADING_COLOR:
                run.font.color.rgb = HEADING_COLOR[lvl]

        elif t == "qa_line":
            # Q:/A: 行：Q加粗，A正常，左缩进
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Cm(0.5)
            text = tok["text"]
            if text.startswith("Q:"):
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(10)
            else:
                run = p.add_run(text)
                run.font.size = Pt(10)

        elif t == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            add_inline(p, tok["text"])

        elif t == "image":
            img_path = resolve_image(tok["src"], images_dir, md_dir)
            if img_path:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(img_path, width=CONTENT_WIDTH)
                    if tok.get("alt"):
                        cap = doc.add_paragraph(tok["alt"])
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if cap.runs:
                            cap.runs[0].italic = True
                            cap.runs[0].font.size = Pt(9)
                            cap.runs[0].font.color.rgb = RGBColor(0x66,0x66,0x66)
                except Exception as e:
                    log(f"  [WARN] 图片插入失败 {img_path}: {e}\n")
                    doc.add_paragraph(f"[图片加载失败: {tok['src']}]")
            else:
                log(f"  [WARN] 图片未找到: {tok['src']}\n")

        elif t == "segment_sep":
            # #### 分段符：原样输出为文本行，单独一段
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            run = p.add_run("####")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        elif t == "tags_raw":
            # tags: @@@tag1@@@tag2@@@ → 原样写入，灰色小字
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(tok["text"])
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        elif t == "tags_line":
            # 旧渲染路径（兼容）
            tags = tok.get("tags", [])
            if tags:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(4)
                run = p.add_run("Tags: " + " | ".join(tags))
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.italic = True

        elif t == "table_image":
            # 表格：若为 .html 文件直接解析为 Word 表格；否则先插PNG再插JSON表格
            from pathlib import Path as _P
            tbl_path = tok["path"]
            base = _P(md_dir) if md_dir else _P(images_dir).parent

            def _find_file(rel_path):
                candidates = [
                    _P(rel_path),
                    base / rel_path,
                    base / _P(rel_path).name,
                ]
                for c in candidates:
                    if c.is_file():
                        return str(c)
                return None

            tbl_lower = tbl_path.lower()
            if tbl_lower.endswith(".html") or tbl_lower.endswith(".htm"):
                # HTML 表格文件 → 直接解析渲染为 Word 表格
                html_path = _find_file(tbl_path)
                if html_path:
                    try:
                        html_content = open(html_path, encoding="utf-8", errors="replace").read()
                        td = parse_html_table(html_content)
                        if td["col_count"]:
                            headers = td["headers"]
                            rows = td["rows"]
                            col_n = td["col_count"]
                            n_rows = (1 if headers else 0) + len(rows)
                            if n_rows > 0:
                                table = doc.add_table(rows=n_rows, cols=col_n)
                                table.style = "Table Grid"
                                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                ri = 0
                                if headers:
                                    for ci, h in enumerate(headers[:col_n]):
                                        cell = table.rows[ri].cells[ci]
                                        cell.text = h
                                        if cell.paragraphs and cell.paragraphs[0].runs:
                                            cell.paragraphs[0].runs[0].bold = True
                                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                        _set_cell_bg(cell, "1F3864")
                                    ri += 1
                                for row_data in rows:
                                    padded = (row_data + [""] * col_n)[:col_n]
                                    for ci, val in enumerate(padded):
                                        table.rows[ri].cells[ci].text = val
                                    ri += 1
                                doc.add_paragraph()
                                log(f"  ✓ HTML表格插入: {tbl_path}\n")
                    except Exception as e:
                        log(f"  [WARN] HTML表格插入失败: {e}\n")
                else:
                    log(f"  [WARN] HTML表格文件未找到: {tbl_path}\n")
            else:
                # ① 插入 PNG 图片
                img_path = _find_file(tbl_path)
                if img_path:
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(img_path, width=CONTENT_WIDTH)
                        log(f"  ✓ 表格PNG插入: {img_path}\n")
                    except Exception as e:
                        log(f"  [WARN] 表格PNG插入失败: {e}\n")
                else:
                    log(f"  [WARN] 表格PNG未找到: {tbl_path}\n")

                # ② 插入真实 Word 表格（JSON格式）
                json_rel = tbl_path.replace(".png", ".json").replace(".PNG", ".json")
                json_path = _find_file(json_rel)
                if json_path:
                    try:
                        import json as _json
                        grid = _json.loads(open(json_path, encoding="utf-8").read())
                        if grid:
                            col_n = max(len(r) for r in grid)
                            n_rows = len(grid)
                            table = doc.add_table(rows=n_rows, cols=col_n)
                            table.style = "Table Grid"
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            for ri, row in enumerate(grid):
                                padded = (row + [""] * col_n)[:col_n]
                                for ci, val in enumerate(padded):
                                    cell = table.rows[ri].cells[ci]
                                    cell.text = val
                                    if ri < 2:
                                        if cell.paragraphs and cell.paragraphs[0].runs:
                                            cell.paragraphs[0].runs[0].bold = True
                                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                        _set_cell_bg(cell, "1F3864")
                            doc.add_paragraph()
                            log(f"  ✓ 真实Word表格插入: {tbl_path}\n")
                    except Exception as e:
                        log(f"  [WARN] Word表格插入失败: {e}\n")
            continue

        elif t == "html_table":
            td = parse_html_table(tok["html"])
            if not td["col_count"]:
                continue
            headers  = td["headers"]
            rows     = td["rows"]
            col_n    = td["col_count"]
            n_rows   = (1 if headers else 0) + len(rows)
            if n_rows == 0:
                continue
            table = doc.add_table(rows=n_rows, cols=col_n)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            ri = 0
            if headers:
                for ci, h in enumerate(headers[:col_n]):
                    cell = table.rows[ri].cells[ci]
                    cell.text = h
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True
                    _set_cell_bg(cell, "D0E4F0")
                ri += 1
            for row_data in rows:
                padded = (row_data + [""] * col_n)[:col_n]
                for ci, val in enumerate(padded):
                    table.rows[ri].cells[ci].text = val
                ri += 1
            doc.add_paragraph()

        elif t == "list_item":
            style = "List Bullet" if not tok["ordered"] else "List Number"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Cm(tok.get("level", 0) * 0.75)
            add_inline(p, tok["text"])

        elif t == "code_block":
            for line in tok["text"].split("\n"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                run = p.add_run(line or " ")
                run.font.name = "Courier New"
                run.font.size = Pt(9)

        elif t == "hr":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)

    doc.save(output_path)
    log(f"✓ 已保存: {output_path}\n")


def run_conversion(md_path: str, images_dir: str, output_path: str, log_cb=None):
    def log(msg):
        if log_cb: log_cb(msg)

    log(f"读取: {md_path}\n")
    with open(md_path, "r", encoding="utf-8", errors="replace") as f:
        md_text = f.read()

    md_dir = str(Path(md_path).parent)
    log("解析 Markdown…\n")
    tokens = parse_markdown(md_text, md_dir=md_dir)
    log(f"  共 {len(tokens)} 个元素\n")

    log("生成 Word 文档…\n")
    build_docx_file(tokens, images_dir, output_path, log_cb=log_cb, md_dir=md_dir)


# ══════════════════════════════════════════════════════════════
#  Tool2 面板
# ══════════════════════════════════════════════════════════════

class Tool2Panel(tk.Frame):
    """MD → Word 转换面板"""

    def __init__(self, parent, shared_state: dict, status_bar,
                 navigate_cb=None, **kw):
        bg = kw.pop("bg", COLORS["bg_main"])
        super().__init__(parent, bg=bg, **kw)
        self._shared     = shared_state
        self._status_bar = status_bar
        self._navigate   = navigate_cb
        self._build()

    def _build(self):
        # 顶部标题
        hdr = tk.Frame(self, bg=COLORS["bg_card"])
        hdr.pack(fill="x")
        if self._navigate:
            tk.Button(hdr, text="◈  主页",
                      bg=COLORS["bg_card"], fg=COLORS["accent"],
                      activebackground=COLORS["bg_hover"],
                      activeforeground=COLORS["accent"],
                      relief="flat", bd=0, cursor="hand2",
                      font=FONTS["sm"], padx=PADDING["md"],
                      command=lambda: self._navigate("home")
                      ).pack(side="right", pady=PADDING["sm"],
                             padx=PADDING["md"])
        tk.Label(hdr, text="⬢  MD → Word", bg=COLORS["bg_card"],
                 fg=COLORS["tool2_color"],
                 font=FONTS["h1"]).pack(side="left",
                                        padx=PADDING["xl"], pady=PADDING["lg"])
        tk.Label(hdr, text="Markdown + 图片文件夹  →  .docx 文档",
                 bg=COLORS["bg_card"], fg=COLORS["text_muted"],
                 font=FONTS["sm"]).pack(side="left", pady=PADDING["lg"])
        Divider(self).pack(fill="x")

        body = tk.Frame(self, bg=COLORS["bg_main"])
        body.pack(fill="both", expand=True, padx=PADDING["xl"],
                  pady=PADDING["xl"])

        # ── STEP 1：MD 文件 ──────────────────────────────────
        step1 = tk.LabelFrame(body, text=" STEP 1  Markdown 文件 ",
                               bg=COLORS["bg_card"],
                               fg=COLORS["tool2_color"],
                               font=FONTS["xs"],
                               bd=1, relief="groove")
        step1.pack(fill="x", pady=(0, PADDING["md"]))

        inner1 = tk.Frame(step1, bg=COLORS["bg_card"])
        inner1.pack(fill="x", padx=PADDING["md"], pady=PADDING["md"])

        self._md_row = FilePickRow(inner1, "MD 文件",
                                   self._browse_md,
                                   bg=COLORS["bg_card"])
        self._md_row.pack(fill="x")

        # 如有共享路径，提示使用
        self._use_shared_btn = StyledButton(
            inner1,
            text="← 使用文档解析的输出结果",
            style="ghost",
            command=self._use_shared_md,
        )
        self._use_shared_btn.pack(anchor="w", pady=(PADDING["sm"], 0))

        # ── STEP 2：图片文件夹 ───────────────────────────────
        step2 = tk.LabelFrame(body, text=" STEP 2  图片文件夹（可选）",
                               bg=COLORS["bg_card"],
                               fg=COLORS["tool2_color"],
                               font=FONTS["xs"],
                               bd=1, relief="groove")
        step2.pack(fill="x", pady=(0, PADDING["md"]))

        inner2 = tk.Frame(step2, bg=COLORS["bg_card"])
        inner2.pack(fill="x", padx=PADDING["md"], pady=PADDING["md"])

        tk.Label(inner2,
                 text="选好 MD 文件后会自动检测同级 images/ 文件夹，也可手动选择",
                 bg=COLORS["bg_card"], fg=COLORS["text_muted"],
                 font=FONTS["xs"],
                 wraplength=700, justify="left").pack(anchor="w",
                                                       pady=(0, PADDING["sm"]))
        self._img_row = FilePickRow(inner2, "images/",
                                    self._browse_img,
                                    placeholder="（自动检测，或手动选择）",
                                    bg=COLORS["bg_card"])
        self._img_row.pack(fill="x")

        # ── STEP 3：输出位置 ─────────────────────────────────
        step3 = tk.LabelFrame(body, text=" STEP 3  输出位置",
                               bg=COLORS["bg_card"],
                               fg=COLORS["tool2_color"],
                               font=FONTS["xs"],
                               bd=1, relief="groove")
        step3.pack(fill="x", pady=(0, PADDING["md"]))

        inner3 = tk.Frame(step3, bg=COLORS["bg_card"])
        inner3.pack(fill="x", padx=PADDING["md"], pady=PADDING["md"])

        self._out_row = FilePickRow(inner3, "保存到",
                                    self._browse_out,
                                    placeholder="（默认：与 MD 文件同目录）",
                                    bg=COLORS["bg_card"])
        self._out_row.pack(fill="x")

        # ── 转换按钮 ─────────────────────────────────────────
        btn_row = tk.Frame(body, bg=COLORS["bg_main"])
        btn_row.pack(pady=PADDING["md"])

        self._convert_btn = StyledButton(btn_row, text="🚀  开始转换",
                                         style="green",
                                         command=self._start)
        self._convert_btn.pack(side="left", padx=(0, PADDING["md"]))

        self._open_dir_btn = StyledButton(btn_row, text="📂  打开输出目录",
                                          style="secondary",
                                          command=self._open_output_dir)
        self._open_dir_btn.pack(side="left")
        self._open_dir_btn.config(state="disabled")

        # ── 进度 ─────────────────────────────────────────────
        self._progress = ProgressRow(body,
                                     bar_style="Green.Horizontal.TProgressbar",
                                     bg=COLORS["bg_main"])
        self._progress.pack(fill="x", pady=(0, PADDING["sm"]))

        # ── 日志 ─────────────────────────────────────────────
        tk.Label(body, text="运行日志",
                 bg=COLORS["bg_main"], fg=COLORS["text_muted"],
                 font=FONTS["xs"]).pack(anchor="w")
        self._log = LogView(body, height=10)
        self._log.pack(fill="both", expand=True, pady=(4, 0))

        self._last_output_dir = None

    # ── 浏览回调 ──────────────────────────────────────────────

    def _browse_md(self):
        p = filedialog.askopenfilename(
            title="选择 Markdown 文件",
            filetypes=[("Markdown 文件", "*.md"), ("所有文件", "*.*")],
        )
        if not p:
            return
        self._md_row.set(p)
        img = Path(p).parent / "images"
        if img.is_dir():
            self._img_row.set(str(img))
        if not self._out_row.get():
            self._out_row.set(str(Path(p).parent))

    def _browse_img(self):
        p = filedialog.askdirectory(title="选择图片文件夹（images/）")
        if p:
            self._img_row.set(p)

    def _browse_out(self):
        p = filedialog.askdirectory(title="选择输出目录")
        if p:
            self._out_row.set(p)

    def _use_shared_md(self):
        md = self._shared.get("last_md_path", "")
        if not md or not Path(md).exists():
            messagebox.showinfo("提示",
                                "当前没有来自「文档解析」的 Markdown 文件。\n"
                                "请先在「文档解析」中完成解析。",
                                parent=self)
            return
        self._md_row.set(md)
        img = self._shared.get("last_images_dir", "")
        if img and Path(img).exists():
            self._img_row.set(img)
        self._out_row.set(str(Path(md).parent))
        self._log.append(f"已加载：{md}", "SUCCESS")

    # ── 转换流程 ──────────────────────────────────────────────

    def _start(self):
        md = self._md_row.get()
        if not md:
            messagebox.showwarning("提示", "请先选择 Markdown 文件", parent=self)
            return
        if not Path(md).is_file():
            messagebox.showerror("错误", f"文件不存在：\n{md}", parent=self)
            return

        img_dir = self._img_row.get()
        if not img_dir:
            auto = Path(md).parent / "images"
            img_dir = str(auto) if auto.is_dir() else str(Path(md).parent)

        out_dir = self._out_row.get() or str(Path(md).parent)
        stem = Path(md).stem
        output = Path(out_dir) / f"{stem}.docx"
        c = 1
        while output.exists():
            output = Path(out_dir) / f"{stem}_{c}.docx"
            c += 1

        self._convert_btn.config(state="disabled")
        self._progress.reset("准备中…")
        self._status_bar.set("转换中…", "running")
        self._log.clear()

        threading.Thread(target=self._worker,
                         args=(md, img_dir, str(output)),
                         daemon=True).start()

    def _worker(self, md, img_dir, output):
        try:
            self._progress.update(20, "解析 Markdown…")
            self._log.append("开始转换…", "INFO")
            run_conversion(md, img_dir, output,
                           log_cb=lambda m: self._log.append(m.strip(), "INFO"))
            self._progress.update(100, "✓ 转换完成")
            self._last_output_dir = str(Path(output).parent)
            self.after(0, self._done_ok, output)
        except Exception:
            err = traceback.format_exc()
            self._log.append(err, "ERROR")
            self._progress.reset("✗ 转换失败")
            self.after(0, self._done_err)

    def _done_ok(self, output):
        self._convert_btn.config(state="normal")
        self._open_dir_btn.config(state="normal")
        self._status_bar.set(f"转换完成 ✓  {Path(output).name}", "success")
        self._log.append(f"已保存至: {output}", "SUCCESS")
        messagebox.showinfo("转换成功 🎉",
                            f"Word 文档已生成：\n{output}",
                            parent=self)

    def _done_err(self):
        self._convert_btn.config(state="normal")
        self._status_bar.set("转换失败", "error")
        messagebox.showerror("转换失败", "请查看下方日志获取详细信息。",
                             parent=self)

    def _open_output_dir(self):
        target = self._last_output_dir
        if not target or not Path(target).exists():
            messagebox.showinfo("提示", "还没有输出结果，请先运行转换。",
                                parent=self)
            return
        import subprocess
        try:
            if sys.platform == "win32":
                os.startfile(target)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", target])
            else:
                subprocess.Popen(["xdg-open", target])
        except Exception as e:
            messagebox.showerror("错误", f"无法打开目录：{e}", parent=self)
